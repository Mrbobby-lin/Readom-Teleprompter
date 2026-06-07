"""
文本管理模块
负责：导入 txt/docx、处理粘贴文本、分句分段
"""

import os
import re


class TextManager:
    """管理提词器中的所有文本"""

    def __init__(self):
        self.paragraphs = []       # 段落列表，每段是一个 Paragraph 对象
        self.sentences = []        # 展平后的句子列表，每句是一个 Sentence 对象
        self.full_text = ""        # 原始全文
        self.title = ""            # 文档标题（文件名或无）
        self.total_chars = 0       # 总字符数

    def load_from_file(self, filepath):
        """从文件加载文本，支持 .txt 和 .docx"""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".txt":
            return self.load_from_txt(filepath)
        elif ext == ".docx":
            return self.load_from_docx(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持 .txt 和 .docx")

    def load_from_txt(self, filepath):
        """加载 .txt 文件"""
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        self.title = os.path.basename(filepath)
        self._set_text(text)
        return True

    def load_from_docx(self, filepath):
        """加载 .docx (Word) 文件"""
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs_text.append(text)
            text = "\n".join(paragraphs_text)
            self.title = os.path.basename(filepath)
            self._set_text(text)
            return True
        except ImportError:
            raise ImportError(
                "请安装 python-docx: pip install python-docx"
            )

    def load_from_text(self, text, title="粘贴文本"):
        """从粘贴的文本加载"""
        self.title = title
        self._set_text(text)
        return True

    def _set_text(self, text):
        """处理原始文本，建立分段和分句索引"""
        self.full_text = text.strip()
        self.paragraphs = []
        self.sentences = []

        # 按换行分段
        raw_paragraphs = re.split(r"\n\s*\n|\r\n\s*\r\n", self.full_text)
        raw_paragraphs = [p.strip() for p in raw_paragraphs if p.strip()]

        for p_idx, para_text in enumerate(raw_paragraphs):
            para = Paragraph(p_idx, para_text)
            # 对段落内分句
            sentence_texts = self._split_sentences(para_text)
            for s_idx, sent_text in enumerate(sentence_texts):
                sent = Sentence(p_idx, s_idx, sent_text)
                para.sentences.append(sent)
                self.sentences.append(sent)
            self.paragraphs.append(para)

        self.total_chars = len(self.full_text.replace(" ", "").replace("\n", ""))

    def _split_sentences(self, text):
        """将文本切分为句子"""
        # 中文分句：句号、问号、感叹号、省略号、分号、冒号
        parts = re.split(r"([。！？…\n])", text)
        sentences = []
        buffer = ""
        for part in parts:
            buffer += part
            if part in "。！？…\n" or part.strip() == "":
                buffer = buffer.strip()
                if buffer:
                    sentences.append(buffer)
                    buffer = ""
        if buffer.strip():
            sentences.append(buffer.strip())
        # 如果没分出来，整段作为一句
        if not sentences:
            sentences = [text.strip()]
        return sentences

    def get_sentence_count(self):
        """获取句子总数"""
        return len(self.sentences)

    def get_paragraph_count(self):
        """获取段落总数"""
        return len(self.paragraphs)

    def has_text(self):
        """是否有文本"""
        return len(self.sentences) > 0

    def get_sentence_text(self, sent_idx):
        """获取指定索引的句子文本"""
        if 0 <= sent_idx < len(self.sentences):
            return self.sentences[sent_idx].text
        return ""

    def get_context(self, sent_idx, window=3):
        """获取某句周围的上下文（用于匹配）"""
        start = max(0, sent_idx - window)
        end = min(len(self.sentences), sent_idx + window + 1)
        return "".join(self.sentences[i].text for i in range(start, end))

    def get_paragraph_sentences(self, p_idx):
        """获取指定段落的所有句子"""
        if 0 <= p_idx < len(self.paragraphs):
            return self.paragraphs[p_idx].sentences
        return []


class Paragraph:
    """段落结构"""
    __slots__ = ("idx", "text", "sentences")

    def __init__(self, idx, text):
        self.idx = idx
        self.text = text
        self.sentences = []


class Sentence:
    """句子结构"""
    __slots__ = ("p_idx", "s_idx", "text", "is_read")

    def __init__(self, p_idx, s_idx, text):
        self.p_idx = p_idx      # 段落索引
        self.s_idx = s_idx      # 句子索引（段内）
        self.text = text        # 句子文本
        self.is_read = False    # 是否已读
